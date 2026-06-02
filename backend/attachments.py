"""
Chat attachments — saving uploaded files to disk, extracting text from
documents, and detecting which models can natively understand images.

A user message can carry attachments (images and documents). Images are passed
to vision-capable LLMs as native multimodal blocks; documents are extracted to
text and injected into the prompt so any model can use them as context.
"""
from __future__ import annotations

import base64
import logging
import os
import re
import uuid

logger = logging.getLogger(__name__)

# Max characters of extracted document text we inject into the prompt.
MAX_DOC_CHARS = 12000

IMAGE_MIME_PREFIX = "image/"
# Document mime types we can extract text from.
TEXT_MIME_TYPES = {
    "text/plain", "text/markdown", "text/csv", "application/json",
    "text/html", "application/xml", "text/xml",
}
PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _uploads_dir(data_dir: str, conversation_id: str) -> str:
    safe_conv = re.sub(r"[^A-Za-z0-9_-]", "_", conversation_id or "misc")
    path = os.path.join(data_dir, "uploads", safe_conv)
    os.makedirs(path, exist_ok=True)
    return path


def _safe_filename(name: str) -> str:
    base = os.path.basename(name or "file")
    return re.sub(r"[^A-Za-z0-9._-]", "_", base)[:120] or "file"


def kind_for(mime_type: str) -> str:
    if mime_type.startswith(IMAGE_MIME_PREFIX):
        return "image"
    return "document"


def decode_b64(b64data: str, mime_type: str = "") -> tuple[bytes, str, str]:
    """Decode a base64 payload that may be a raw string or a `data:` URI.

    Returns (raw_bytes, clean_base64, mime_type).
    """
    if b64data.startswith("data:"):
        # data:<mime>;base64,<payload>
        try:
            header, b64data = b64data.split(",", 1)
            if not mime_type and ":" in header and ";" in header:
                mime_type = header.split(":", 1)[1].split(";", 1)[0]
        except ValueError:
            pass
    return base64.b64decode(b64data), b64data, mime_type


def extract_text_from_bytes(raw: bytes, mime_type: str) -> str:
    """Best-effort text extraction from in-memory bytes (no disk write)."""
    import io
    try:
        if mime_type in TEXT_MIME_TYPES or mime_type.startswith("text/"):
            return raw.decode("utf-8", errors="replace")[: MAX_DOC_CHARS * 2]
        if mime_type == PDF_MIME:
            try:
                from pypdf import PdfReader
            except ImportError:
                return ""
            reader = PdfReader(io.BytesIO(raw))
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
                if sum(len(p) for p in parts) > MAX_DOC_CHARS * 2:
                    break
            return "\n".join(parts)
        if mime_type == DOCX_MIME:
            try:
                import docx
            except ImportError:
                return ""
            return "\n".join(p.text for p in docx.Document(io.BytesIO(raw)).paragraphs)
    except Exception as e:
        logger.error("Failed to extract text from bytes: %s", e)
    return ""


def save_attachment(
    data_dir: str,
    conversation_id: str,
    name: str,
    mime_type: str,
    b64data: str,
) -> dict:
    """Persist an uploaded file to disk and return its metadata.

    `b64data` may be a raw base64 string or a `data:` URI.
    """
    raw, b64data, mime_type = decode_b64(b64data, mime_type)
    stored_name = f"{uuid.uuid4().hex[:8]}_{_safe_filename(name)}"
    dest_dir = _uploads_dir(data_dir, conversation_id)
    dest_path = os.path.join(dest_dir, stored_name)
    with open(dest_path, "wb") as f:
        f.write(raw)

    safe_conv = os.path.basename(dest_dir)
    return {
        "id": uuid.uuid4().hex,
        "name": name,
        "mime_type": mime_type or "application/octet-stream",
        "kind": kind_for(mime_type or ""),
        "filename": stored_name,
        "url": f"/api/uploads/{safe_conv}/{stored_name}",
        "size": len(raw),
        "path": dest_path,
    }


def extract_text(path: str, mime_type: str) -> str:
    """Best-effort text extraction from a document. Returns '' on failure."""
    try:
        if mime_type in TEXT_MIME_TYPES or mime_type.startswith("text/"):
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return f.read(MAX_DOC_CHARS * 2)

        if mime_type == PDF_MIME:
            try:
                from pypdf import PdfReader
            except ImportError:
                logger.warning("pypdf not installed — cannot extract PDF text")
                return ""
            reader = PdfReader(path)
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
                if sum(len(p) for p in parts) > MAX_DOC_CHARS * 2:
                    break
            return "\n".join(parts)

        if mime_type == DOCX_MIME:
            try:
                import docx  # python-docx
            except ImportError:
                logger.warning("python-docx not installed — cannot extract DOCX text")
                return ""
            d = docx.Document(path)
            return "\n".join(p.text for p in d.paragraphs)
    except Exception as e:
        logger.error("Failed to extract text from %s: %s", path, e)
    return ""


def read_base64(path: str) -> str:
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode("ascii")


# ── Vision capability detection ──────────────────────────────────────────────

def model_supports_vision(provider: str, model: str) -> bool:
    """Heuristic: does this provider+model understand images natively?"""
    p = (provider or "").lower()
    m = (model or "").lower()

    # Generic signals present across providers.
    if any(tok in m for tok in ("vision", "-vl", "vl-", "llava", "minicpm-v", "4v")):
        return True

    if p == "anthropic":
        # Claude 3 and later are all multimodal.
        return "claude" in m

    if p in ("openai", "openrouter", "custom"):
        return any(tok in m for tok in ("gpt-4o", "gpt-4.1", "gpt-4-turbo", "gpt-5", "o1", "o3", "o4", "gemini", "claude"))

    if p == "gemini":
        # All Gemini 1.5+ are multimodal.
        return "gemini" in m

    if p == "ollama":
        return any(tok in m for tok in ("gemma3", "llama3.2-vision", "qwen2.5vl", "qwen2-vl", "minicpm", "moondream", "bakllava"))

    if p == "glm":
        return "4v" in m

    return False


# Preferred Ollama Cloud vision models, in order of preference. Used to decode
# images on behalf of a non-vision text model so the user doesn't have to swap
# their model just to attach a picture.
OLLAMA_CLOUD_VISION_MODELS = ["gemma3:27b-cloud", "qwen2.5vl:32b-cloud", "llama3.2-vision:90b"]


async def describe_images(adapter, image_parts: list, model: str, prompt_hint: str = "") -> str:
    """Ask a vision-capable model to describe images, returning text.

    `image_parts` are normalized image content parts. `adapter` is any adapter
    whose `model` supports vision (e.g. an Ollama Cloud adapter + vision model).
    Returns '' on failure.
    """
    if not image_parts:
        return ""
    system = (
        "You are a precise visual analyst. Describe the attached image(s) in "
        "thorough, factual detail: text/labels, objects, people, layout, colours, "
        "charts/numbers, and anything relevant. Be objective and complete."
    )
    user_text = "Describe the attached image(s) in detail."
    if prompt_hint:
        user_text += f"\n\nThe user's request, for context: {prompt_hint}"
    content = [{"type": "text", "text": user_text}, *image_parts]
    try:
        resp = await adapter.complete(
            system_prompt=system,
            messages=[{"role": "user", "content": content}],
            model=model,
            max_tokens=1024,
        )
        return resp.content or ""
    except Exception as e:
        logger.error("Vision description failed (model=%s): %s", model, e)
        return ""


def vision_unsupported_message(provider: str, model: str) -> str:
    """Friendly guidance shown when an image is attached but the model can't see it."""
    return (
        f"⚠️ The current model (**{provider} / {model}**) can't read images. "
        "I've used any attached documents and your text, but to analyse images "
        "please configure a vision-capable model — for example Claude (Anthropic), "
        "GPT-4o (OpenAI), Gemini, or an Ollama Cloud vision model such as "
        "`gemma3:27b-cloud` or `llama3.2-vision` (add the provider's API key under "
        "**Settings → API**, then set it on the CEO agent)."
    )
